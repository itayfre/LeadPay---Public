"""
Word renderer — converts a report payload dict into .docx bytes using
python-docx. RTL paragraph and run properties follow the hebrew-document-
generator skill conventions (w:bidi on paragraph, w:rtl on run).
Default font is David (system-safe, formal Hebrew).
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


_FONT = "David"
_BODY_PT = 11
_HEAD_PT = 16
_SUB_PT = 13


# ─── RTL helpers (from hebrew-document-generator skill) ──────────────────────

def _set_paragraph_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.makeelement(qn("w:bidi"), {})
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _set_run_rtl(run, size: int = _BODY_PT) -> None:
    rPr = run._r.get_or_add_rPr()
    rtl = rPr.makeelement(qn("w:rtl"), {})
    rPr.append(rtl)
    run.font.name = _FONT
    run.font.size = Pt(size)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), _FONT)
    rFonts.set(qn("w:ascii"), _FONT)
    rFonts.set(qn("w:hAnsi"), _FONT)


# ─── Document helpers ─────────────────────────────────────────────────────────

def _add_para(doc: Document, text: str, bold: bool = False, size: int = _BODY_PT):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    _set_run_rtl(r, size=size)
    _set_paragraph_rtl(p)
    return p


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[Any]],
    totals_row: list[Any] | None = None,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        para = hdr_cells[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        _set_run_rtl(run)
        _set_paragraph_rtl(para)

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            para = cells[i].paragraphs[0]
            run = para.add_run(str(val))
            _set_run_rtl(run)
            _set_paragraph_rtl(para)

    if totals_row is not None:
        cells = table.add_row().cells
        for i, val in enumerate(totals_row):
            para = cells[i].paragraphs[0]
            run = para.add_run(str(val))
            run.bold = True
            _set_run_rtl(run)
            _set_paragraph_rtl(para)


def _shekel(n: float) -> str:
    if n is None:
        return "—"
    return f"₪{n:,.0f}"


# ─── Public API ───────────────────────────────────────────────────────────────

def render_report_docx(payload: dict) -> bytes:
    doc = Document()

    # Set default style
    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(_BODY_PT)

    # Title
    _add_para(doc, "דוח הכנסות והוצאות", bold=True, size=_HEAD_PT)

    # Building metadata
    b = payload["building"]
    _add_para(doc, f"{b['name']} · {b['address']}, {b['city']}")
    _add_para(doc, payload["period"]["label"])

    # Summary
    doc.add_paragraph()
    _add_para(doc, "סיכום", bold=True, size=_SUB_PT)
    s = payload["summary"]
    _add_table(
        doc,
        ["מאזן נוכחי", "סה״כ הוצאות", "סה״כ הכנסות"],
        [[_shekel(s["net_balance"]), _shekel(s["total_expenses"]), _shekel(s["total_income"])]],
    )

    # Income table
    doc.add_paragraph()
    _add_para(doc, "פירוט הכנסות לפי דייר", bold=True, size=_SUB_PT)
    if b.get("expected_monthly_payment"):
        _add_para(doc, f"דמי ועד חודשי: {_shekel(b['expected_monthly_payment'])} לדירה")

    cols = payload["period"]["columns"]
    income_headers = ["דירה", "שם דייר"] + [c["label"] for c in cols] + ["שולם", "לתשלום", "יתרה"]
    income_rows = []
    for r in payload["income_by_tenant"]:
        row = (
            [str(r["apartment_number"]), r["tenant_name"]]
            + [_shekel(c["amount"]) for c in r["cells"]]
            + [_shekel(r["paid_total"]), _shekel(r["expected_total"]),
               _shekel(r["balance"]) if r["balance"] > 0 else "—"]
        )
        income_rows.append(row)
    tot = payload["income_totals_row"]
    income_totals = (
        ["סה״כ", ""]
        + [_shekel(c["amount"]) for c in tot["cells"]]
        + [_shekel(tot["paid_total"]), _shekel(tot["expected_total"]),
           _shekel(tot["balance"]) if tot["balance"] > 0 else "—"]
    )
    _add_table(doc, income_headers, income_rows, totals_row=income_totals)

    # Expenses table
    doc.add_paragraph()
    _add_para(doc, "פירוט הוצאות", bold=True, size=_SUB_PT)
    if not payload["expenses_by_month"]:
        _add_para(doc, "אין הוצאות בתקופה זו")
    else:
        exp_rows = []
        for g in payload["expenses_by_month"]:
            for i, row in enumerate(g["rows"]):
                exp_rows.append([g["month_label"] if i == 0 else "",
                                 row["description"], row["category"],
                                 _shekel(row["amount"])])
            exp_rows.append([f"סה״כ {g['month_label']}", "", "",
                              _shekel(g["subtotal"])])
        exp_rows.append(["סה״כ הוצאות", "", "",
                         _shekel(payload["expenses_grand_total"])])
        _add_table(doc, ["חודש", "תיאור", "קטגוריה", "סכום"], exp_rows)

    # Debtors
    if payload["debtors_period"] or payload["debtors_lifetime"]:
        doc.add_paragraph()
        _add_para(doc, "חייבים – יתרת חוב פתוח", bold=True, size=_SUB_PT)
        for title, debtors in [
            ("חוב לתקופה זו", payload["debtors_period"]),
            ("יתרת חוב כוללת", payload["debtors_lifetime"]),
        ]:
            if not debtors:
                continue
            _add_para(doc, title, bold=True)
            _add_table(
                doc,
                ["דירה", "שם דייר", "חוב", "הערה"],
                [[str(d["apartment_number"]), d["tenant_name"],
                  _shekel(d["debt"]), d["note"]] for d in debtors],
            )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_tenant_report_docx(payload: dict) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(_BODY_PT)

    _add_para(doc, payload["tenant"]["name"], bold=True, size=_HEAD_PT)
    t = payload["tenant"]
    meta = f"דירה {t['apartment_number']}"
    if t.get("floor"):
        meta += f" · קומה {t['floor']}"
    meta += f" · {t['building']['name']}, {t['building']['address']}, {t['building']['city']}"
    _add_para(doc, meta)
    _add_para(doc, payload["period"]["label"], bold=True)
    if t.get("standing_order"):
        so = t["standing_order"]
        line = "הוראת קבע פעילה"
        if so.get("amount"):
            line += f" — {_shekel(so['amount'])} לחודש"
        if so.get("end_date"):
            line += f" (עד {so['end_date']})"
        _add_para(doc, line)

    doc.add_paragraph()
    _add_para(doc, "סיכום", bold=True, size=_SUB_PT)
    s = payload["summary"]
    _add_table(
        doc,
        ["סה״כ שולם בתקופה", "חוב כולל", "חוב לתקופה"],
        [[_shekel(s["period_paid"]), _shekel(s["lifetime_debt"]), _shekel(s["period_debt"])]],
    )

    doc.add_paragraph()
    _add_para(doc, "פירוט חודשי", bold=True, size=_SUB_PT)
    rows = []
    for r in payload["months"]:
        status_he = {"paid": "שולם", "partial": "חלקי", "unpaid": "לא שולם"}.get(r["status"], "")
        diff = r["difference"]
        diff_s = ("-" + _shekel(-diff)) if diff < 0 else _shekel(diff)
        rows.append([r["period_label"], _shekel(r["expected"]), _shekel(r["paid"]), diff_s, status_he])
    _add_table(doc, ["חודש", "צפוי", "שולם", "הפרש", "סטטוס"], rows)

    doc.add_paragraph()
    _add_para(doc, "תנועות בתקופה", bold=True, size=_SUB_PT)
    if not payload["transactions"]:
        _add_para(doc, "אין תנועות בתקופה זו")
    else:
        tx_rows = []
        for tx in payload["transactions"]:
            date_str = tx["date"].split("T")[0] if "T" in tx["date"] else tx["date"]
            tx_rows.append([date_str, tx["description"], _shekel(tx["amount"]), tx["method"]])
        _add_table(doc, ["תאריך", "תיאור", "סכום", "אופן תשלום"], tx_rows)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
